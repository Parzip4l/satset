from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("UAT_QA_Manual_Fitur_GA.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_w = row.cells[idx]._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                row.cells[idx]._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def keep_table_rows_together(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header_row(table):
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def style_table(table, widths):
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_width(table, widths)
    keep_table_rows_together(table)
    repeat_header_row(table)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(31, 77, 120)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181) if level < 3 else RGBColor(31, 77, 120)
    return p


def add_meta_table(doc):
    rows = [
        ("Versi", "1.0"),
        ("Tanggal", "20 Juli 2026"),
        ("Aplikasi", "SATSET LRTJ"),
        ("Modul", "General Affairs / BUM"),
        ("Tujuan", "Panduan UAT oleh user bisnis dan QA manual untuk seluruh fitur GA/BUM."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Keterangan"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    style_table(table, [1.55, 4.95])


def add_scope(doc):
    add_heading(doc, "1. Ruang Lingkup", 1)
    items = [
        "Login, dashboard, menu tiket, ticket umum GA.",
        "Permintaan konsumsi rapat: submit, approval atasan, update flow BUM, upload bukti.",
        "Permintaan ATK/RTK: threshold approval, review BUM, procurement state, handover, stok keluar.",
        "GA Permintaan dan Temuan: input laporan, evidence, monitoring dashboard.",
        "Inventory BUM: master barang, detail barang, adjustment, kartu stok, receiving, opname, laporan.",
        "Analytics BUM, master data pendukung, notifikasi, booking ruang meeting, vCard, dan API.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_roles(doc):
    add_heading(doc, "2. Role dan Data Uji", 1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Role"
    table.rows[0].cells[1].text = "Fokus Pengujian"
    rows = [
        ("Requester/Karyawan", "Membuat permintaan, melihat ticket sendiri, upload bukti, komentar."),
        ("Atasan/Approver", "Approve/reject konsumsi dan ATK/RTK yang membutuhkan approval."),
        ("Admin/BUM", "Review ticket GA/BUM, update flow, handover, kelola stok, receiving, opname, laporan."),
        ("Admin Sistem", "Kelola master data, user, role, status, priority, impact, urgency, schema."),
    ]
    for role, focus in rows:
        cells = table.add_row().cells
        cells[0].text = role
        cells[1].text = focus
    style_table(table, [1.75, 4.75])

    p = doc.add_paragraph()
    p.add_run("Data uji minimum: ").bold = True
    p.add_run("user requester, approver, admin/BUM; status Open/In Progress/Pending/Resolved/Closed; priority/impact/urgency Medium; category BUM, konsumsi, ATK/RTK, GA Permintaan & Temuan; minimal tiga item consumable aktif dan satu ruang meeting.")


def add_criteria(doc):
    add_heading(doc, "3. Kriteria Pass/Fail", 1)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Status"
    table.rows[0].cells[1].text = "Kriteria"
    rows = [
        ("Pass", "Hasil aktual sesuai expected result, data tersimpan, history tercatat, workflow benar, stok berubah sesuai transaksi, validasi berjalan, dan tidak ada error 500/UI blocker."),
        ("Fail", "Data tidak tersimpan, status salah, stok minus, role mengakses data yang tidak semestinya, upload tidak tervalidasi, filter/sort gagal, atau API tidak sesuai kontrak."),
    ]
    for status, criteria in rows:
        cells = table.add_row().cells
        cells[0].text = status
        cells[1].text = criteria
    style_table(table, [1.1, 5.4])


TEST_CASES = [
    ("UAT-GA-01", "Login dan Dashboard", "Login valid lalu buka Dashboard, Ticket, BUM Dashboard.", "User masuk, menu sesuai role, halaman tampil tanpa error."),
    ("UAT-GA-02", "Ticket Umum GA", "Buat ticket umum dengan title, description, category, priority/impact/urgency.", "Ticket no TCK-{KODE}-{MMDDYY}-{RUNNING}, status Open, history create tercatat."),
    ("UAT-GA-03", "List/Search/Filter Ticket", "Cari ticket no/title/requester; filter status, priority, category, department.", "Hasil sesuai filter; non-admin hanya melihat ticket sendiri."),
    ("UAT-GA-04", "Detail Ticket", "Assign user/departemen, update status, tambah komentar.", "Assignment, status, history, dan komentar tersimpan."),
    ("UAT-GA-05", "Permintaan Konsumsi", "Isi kegiatan, tipe, tanggal/jam, lokasi, peserta, konsumsi, alasan, approver, upload opsional.", "request_type consumption; WAITING_MANAGER_APPROVAL; lampiran valid tersimpan."),
    ("UAT-GA-06", "Approval Konsumsi", "Approver approve dan reject pada data berbeda.", "Approve ke WAITING_BUM_VERIFICATION; reject ke REJECTED_BY_MANAGER; self-approval ditolak."),
    ("UAT-GA-07", "Flow BUM Konsumsi", "Update APPROVED_BY_BUM sampai CLOSED dengan vendor, tanggal, biaya, catatan.", "Payload dan history berubah sesuai setiap tahap."),
    ("UAT-GA-08", "Bukti Konsumsi", "Upload attendance/documentation/report/material; coba file invalid.", "File valid tersimpan dengan tipe benar; invalid ditolak."),
    ("UAT-GA-09", "ATK/RTK Di Bawah Threshold", "Submit ATK/RTK total di bawah threshold.", "WAITING_BUM_REVIEW; approver tidak wajib; total terhitung."),
    ("UAT-GA-10", "ATK/RTK Di Atas Threshold", "Submit tanpa approver lalu dengan approver.", "Tanpa approver ditolak; dengan approver WAITING_MANAGER_APPROVAL."),
    ("UAT-GA-11", "Approval ATK/RTK", "Approve/reject ticket ATK/RTK yang butuh approval.", "Approve ke WAITING_BUM_REVIEW; reject ke REJECTED_BY_MANAGER."),
    ("UAT-GA-12", "Review BUM ATK/RTK", "Set STOCK_CHECKED, WAITING_PROCUREMENT, READY_TO_HANDOVER, CANCELLED.", "workflow_status, approved_qty, notes, dan history tersimpan."),
    ("UAT-GA-13", "Handover ATK/RTK", "Pilih item, fulfilled_qty, penerima; coba qty melebihi approval/stok.", "Stok OUT tercatat; HANDED_OVER; validasi qty/stok berjalan."),
    ("UAT-GA-14", "Gudang ATK/RTK", "Cek statistik waiting/procurement/ready/handover, active tickets, low stock.", "Angka dan list sesuai data; pagination/link detail bekerja."),
    ("UAT-GA-15", "GA Permintaan & Temuan", "Submit tipe Permintaan dan Temuan dengan lokasi, deskripsi, expected action, evidence.", "WAITING_BUM_REVIEW; title otomatis; evidence valid tersimpan."),
    ("UAT-GA-16", "Dashboard BUM", "Buka dashboard; filter GA all/Permintaan/Temuan.", "Kartu, breakdown, recent ticket menyesuaikan filter."),
    ("UAT-GA-17", "Master Barang", "Tambah/edit item, aktif/nonaktif, filter category, search, sort.", "Code unik; initial stock membuat movement IN; list sesuai filter."),
    ("UAT-GA-18", "Detail/Adjustment Stok", "Adjustment IN/OUT; coba OUT melebihi stok.", "Movement balance benar; stok minus ditolak; trend tampil."),
    ("UAT-GA-19", "Kartu Stok", "Filter item/tanggal dan sort kolom.", "Mutasi sesuai filter; balance_before/after benar."),
    ("UAT-GA-20", "Penerimaan Pengadaan", "Buat receiving dengan vendor, PO/DO/GR, tanggal, item lines.", "Reference RCV-{YYYYMMDD-HHMMSS}; status SUBMITTED; lines tersimpan."),
    ("UAT-GA-21", "Update Receiving", "Set RECEIVED/PARTIALLY_RECEIVED/REJECTED/STORED/CLOSED dengan qty.", "Delta qty_received menambah stok; tidak menggandakan stok lama."),
    ("UAT-GA-22", "Stock Opname", "Buat opname period YYYY-MM dengan physical_stock.", "Status CLOSED; variance dihitung; ADJUSTMENT tercatat jika beda."),
    ("UAT-GA-23", "Laporan Bulanan", "Pilih periode laporan.", "Usage OUT, receiving, dan konsumsi sesuai bulan."),
    ("UAT-GA-24", "Analytics BUM", "Buka analytics dan semua endpoint chart.", "JSON success true; chart/data render tanpa error."),
    ("UAT-GA-25", "Master Data Ticket", "Kelola division, lokasi, PIC, department, category, priority, status, impact, urgency, schema.", "CRUD/list berjalan; schema category bisa diambil."),
    ("UAT-GA-26", "Notifikasi", "Mark read satu item, mark all, clear all.", "Read state/counter/list berubah sesuai aksi."),
    ("UAT-GA-27", "Booking Ruang Meeting", "Kelola room, buat booking calendar, cek events, hapus booking.", "Booking tampil; event API valid; delete berhasil."),
    ("UAT-GA-28", "VCard/Kontak", "Buka contact dan download vCard.", "Profil tampil; file vCard terunduh."),
    ("UAT-GA-29", "API Master/Ticket", "Panggil /api/v1/master/*, tickets, history, POST/PUT, schema.", "JSON valid; validasi input berjalan; data sesuai."),
    ("UAT-GA-30", "API Analytics", "Panggil /api/bum/analytics/* dan /api/v1/bum/analytics/*.", "success true dan data; filter invalid tidak error 500."),
]


def add_test_cases(doc):
    add_heading(doc, "4. Skenario UAT End-to-End", 1)
    table = doc.add_table(rows=1, cols=4)
    headers = ["ID", "Fitur", "Langkah Ringkas", "Expected Result"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for case in TEST_CASES:
        cells = table.add_row().cells
        for idx, value in enumerate(case):
            cells[idx].text = value
    style_table(table, [0.8, 1.35, 2.25, 2.1])


def add_checklist(doc):
    add_heading(doc, "5. Checklist QA Manual Per Build", 1)
    checks = [
        "Login valid/invalid diuji.",
        "Role requester tidak dapat melihat ticket user lain.",
        "Admin/BUM dapat melakukan aksi review dan inventory.",
        "Semua form wajib menolak field kosong yang required.",
        "Upload menerima format yang diizinkan dan menolak format/ukuran tidak valid.",
        "Status/workflow ticket berubah sesuai alur.",
        "History tercatat untuk create, status update, assign, approval, review, handover, flow konsumsi.",
        "Email failure tidak membuat proses utama gagal.",
        "Stok tidak pernah menjadi negatif.",
        "Stock card balance_before/balance_after sesuai transaksi.",
        "Search, filter, sort, pagination bekerja.",
        "Halaman responsive minimal desktop dan mobile.",
        "API mengembalikan JSON valid dan tidak membocorkan stack trace.",
    ]
    for check in checks:
        doc.add_paragraph(check, style="List Bullet")


def add_defects_and_signoff(doc):
    add_heading(doc, "6. Template Log Defect", 1)
    table = doc.add_table(rows=2, cols=9)
    headers = ["Defect ID", "Modul", "Skenario", "Severity", "Step", "Expected", "Actual", "Status", "Owner"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for idx, value in enumerate(["DF-001", "", "", "Critical/High/Medium/Low", "", "", "", "Open", ""]):
        table.rows[1].cells[idx].text = value
    style_table(table, [0.65, 0.7, 0.85, 0.85, 0.9, 0.9, 0.85, 0.7, 0.5])

    add_heading(doc, "7. Sign-Off UAT", 1)
    table = doc.add_table(rows=4, cols=6)
    headers = ["Nama", "Role", "Keputusan", "Catatan", "Tanggal", "Tanda Tangan"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    rows = [
        ["", "User Bisnis/GA", "Accept / Accept with Notes / Reject", "", "", ""],
        ["", "QA", "Pass / Pass with Defects / Fail", "", "", ""],
        ["", "IT/Developer", "Release / Hold", "", "", ""],
    ]
    for ridx, row in enumerate(rows, start=1):
        for cidx, value in enumerate(row):
            table.rows[ridx].cells[cidx].text = value
    style_table(table, [0.95, 1.1, 1.65, 1.25, 0.75, 0.8])


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in [
        ("Heading 1", 16, 18, 10, RGBColor(46, 116, 181)),
        ("Heading 2", 13, 14, 7, RGBColor(46, 116, 181)),
        ("Heading 3", 12, 10, 5, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Dokumen UAT dan QA Manual Fitur GA")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.add_run("SATSET LRTJ - General Affairs / BUM").bold = True

    add_meta_table(doc)
    add_scope(doc)
    add_roles(doc)
    add_criteria(doc)
    add_test_cases(doc)
    add_checklist(doc)
    add_defects_and_signoff(doc)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run("UAT & QA Manual GA | SATSET LRTJ")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
